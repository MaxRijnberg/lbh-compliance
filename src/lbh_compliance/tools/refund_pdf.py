from typing import Literal, Union, TypedDict, List, Optional, Tuple
from pathlib import Path
from PyPDF2 import PdfReader
from docx import Document

import re

from lbh_compliance.config.patterns import PORTCALL_ID_PATTERN, SWIFT_PATTERN
from lbh_compliance.utils.loggers import get_logger


class RefundData(TypedDict):
    beneficiary_account_name: str
    beneficiary_banks: List[Tuple[str, str]]  # Tuple = (NAME, SWIFT)
    intermediary_banks: Optional[List[Tuple[str, str]]]  # Tuple = (NAME, SWIFT)
    portcall_ids: Optional[List[str]]


class ChineseRefundData(TypedDict):
    vessels: List[Tuple[str, str]]  # Tuple = (NAME, IMO)
    banks: List[str]
    account_names: List[str]


class RefundReader:
    def __init__(
        self,
        filename: Union[str, Path],
        logging_level: Literal["debug", "info", "warn", "error", "critical"] = "info",
    ) -> None:
        self.logger = get_logger(__name__, logging_level)
        self.logging_level: Literal["debug", "info", "warn", "error", "critical"] = (
            logging_level
        )

        self.filename: Union[str, Path] = filename

        # initialise file data
        self.file_text: str = ""

    def get_data(self) -> RefundData:
        """
        Retrieve the data from the Refund Request form.

        Returns:
            RefundData: A Dict-like object with the relevant data of the refund request
        """
        self.file_text = self._read_file(self.filename)
        data = self._parse_text(self.file_text)
        return data

    def _read_file(self, filename: Union[str, Path]) -> str:
        """
        Reads the Refund Request file.

        Args:
            filename (Union[str, Path]): The filename to read.

        Raises:
            ValueError: If the filetype is not .pdf or .docx

        Returns:
            str: The text of the file
        """
        ext = str(filename).lower()

        if ext.endswith(".pdf"):
            return self._read_pdf(filename)
        elif ext.endswith(".docx"):
            return self._read_docx(filename)
        else:
            msg = f"Unsupported file type: {ext}"
            self.logger.error(msg)
            raise ValueError(msg)

    def _parse_text(self, txt: str) -> RefundData:
        """
        Filters and extracts the text for relevant data in the screening process.

        Args:
            txt (str): The text to be checked.

        Returns:
            RefundData: A Dict-like object containing relevant information.
        """
        txt = self._normalise(txt)

        beneficiary_name = self._extract_field(txt, "Beneficiary Account Name")

        beneficiary_bank = self._extract_field(txt, "Beneficiary Bank Name")
        beneficiary_swift = self._extract_swift_near(txt, "Beneficiary Bank Swift Code")

        beneficiary_banks = []
        if beneficiary_bank and beneficiary_swift:
            beneficiary_banks.append((beneficiary_bank, beneficiary_swift))

        intermediary_banks = self._extract_intermediary_banks(txt)

        portcalls = re.findall(PORTCALL_ID_PATTERN, txt)
        portcalls = list(set(portcalls)) if portcalls else None

        return {
            "beneficiary_account_name": beneficiary_name,
            "beneficiary_banks": beneficiary_banks,
            "intermediary_banks": intermediary_banks,
            "portcall_ids": portcalls,
        }

    def _read_pdf(self, filename: Union[str, Path]) -> str:
        """
        Reads a PDF file and outputs its text in a str

        Args:
            filename (Union[str, Path]): The filename to read.

        Returns:
            str: The text found in the file
        """
        reader = PdfReader(filename)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or "")
        return "\n".join(text)

    def _read_docx(self, filename: Union[str, Path]) -> str:
        """
        Reads a MS Word/.docx file and outputs its text in a str

        Args:
            filename (Union[str, Path]): The filename to read.

        Returns:
            str: The text found in the file
        """
        doc = Document(str(filename))
        text_parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                text_parts.append(p.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = [
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                ]
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n".join(text_parts)

    def _normalise(self, txt: str) -> str:
        """
        Normalises any text, by replacing:
        - Any space -> " "
        - "–" -> "-" (U+2013 -> U+002d)
        - "|" -> " "

        Args:
            txt (str): The text to clean

        Returns:
            str: The clean text
        """
        txt = re.sub(r"\s+", " ", txt)
        txt = txt.replace("–", "-")
        txt = txt.replace("|", " ")
        return txt

    def _extract_field(self, txt: str, field_name: str) -> str:
        """
        Extracts the text from a field if found in the table.

        Args:
            txt (str): The text to search in.
            field_name (str): The field to search for.

        Returns:
            str: The value of that field, if found, else ""
        """
        pattern = re.compile(rf"{field_name}\s*(.*?)\s*(?:\d+\.|$)", re.IGNORECASE)
        match = pattern.search(txt)
        return match.group(1).strip() if match else ""

    def _extract_swift_near(self, txt: str, label: str) -> str:
        """
        Searches for the field in the table, and ecxtracts the SWIFT code

        Args:
            txt (str): The raw text to search in
            label (str): The value in the row to match

        Returns:
            str: The SWIFT code if found, else ""
        """
        pattern = rf"{label}\s*(.*?)\s*(?:\d+\.|$)"
        match = re.search(pattern, txt, re.IGNORECASE)

        if not match:
            return ""

        segment = match.group(1)

        swift_match = re.search(SWIFT_PATTERN, segment)
        return swift_match.group(0) if swift_match else ""

    def _extract_intermediary_banks(self, txt: str) -> Optional[List[Tuple[str, str]]]:
        """
        Extracts the banks of the intermediary banks, if found.

        Args:
            txt (str): The text to search the intermediary banks in

        Returns:
            Optional[List[Tuple[str, str]]]: A list of Tuples containing (str, str),
                in which the tuples correspond to (bank_name, swift_code)
        """
        banks = []

        name_patterns = [
            "Correspondent Bank Name",
            "Intermediary Bank Name",
        ]

        for name_label in name_patterns:
            name = self._extract_field(txt, name_label)

            if not name:
                continue

            swift = self._extract_swift_near(txt, "Intermediary Bank Swift Code")

            if name:
                banks.append((name, swift or ""))

        return banks or None
